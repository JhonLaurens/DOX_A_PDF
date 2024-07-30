import sys
import os
from PyQt6.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QPushButton, 
                             QFileDialog, QLabel, QProgressBar, QWidget, QComboBox, QMessageBox)
from PyQt6.QtGui import QFont
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from docx import Document
from docx2pdf import convert
import pandas as pd

class ConversionThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, input_file, output_format):
        super().__init__()
        self.input_file = input_file
        self.output_format = output_format

    def run(self):
        try:
            if self.output_format == 'PDF':
                self.convert_to_pdf()
            elif self.output_format == 'CSV':
                self.convert_to_csv()
            elif self.output_format == 'XLSX':
                self.convert_to_xlsx()
        except Exception as e:
            self.error.emit(str(e))

    def convert_to_pdf(self):
        try:
            pdf_output_path = os.path.splitext(self.input_file)[0] + ".pdf"
            convert(self.input_file, pdf_output_path)
            self.progress.emit(100)
            self.finished.emit(pdf_output_path)
        except Exception as e:
            self.error.emit(f"PDF conversion error: {str(e)}")

    def convert_to_csv(self):
        try:
            doc = Document(self.input_file)
            data = [para.text for para in doc.paragraphs]
            df = pd.DataFrame(data, columns=['Content'])
            
            csv_output_path = os.path.splitext(self.input_file)[0] + ".csv"
            df.to_csv(csv_output_path, index=False)
            
            self.finished.emit(csv_output_path)
        except Exception as e:
            self.error.emit(f"CSV conversion error: {str(e)}")

    def convert_to_xlsx(self):
        try:
            doc = Document(self.input_file)
            data = [para.text for para in doc.paragraphs]
            df = pd.DataFrame(data, columns=['Content'])
            
            xlsx_output_path = os.path.splitext(self.input_file)[0] + ".xlsx"
            df.to_excel(xlsx_output_path, index=False)
            
            self.finished.emit(xlsx_output_path)
        except Exception as e:
            self.error.emit(f"XLSX conversion error: {str(e)}")

class DocxConverter(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle("DOCX Converter")
        self.setGeometry(300, 300, 400, 300)
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f0f0;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px 20px;
                margin: 10px 0px;
                border-radius: 5px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QLabel {
                font-size: 14px;
                margin: 10px 0px;
            }
            QProgressBar {
                border: 2px solid grey;
                border-radius: 5px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                width: 10px;
                margin: 0.5px;
            }
        """)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        self.label = QLabel("Select a DOCX file to convert", self)
        layout.addWidget(self.label)

        self.format_combo = QComboBox(self)
        self.format_combo.addItems(['PDF', 'CSV', 'XLSX'])
        layout.addWidget(self.format_combo)

        self.btnOpenFile = QPushButton("Open DOCX file", self)
        self.btnOpenFile.clicked.connect(self.openFileDialog)
        layout.addWidget(self.btnOpenFile)

        self.progress_bar = QProgressBar(self)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

    def openFileDialog(self):
        try:
            options = QFileDialog.Option.ReadOnly
            
            filePath, _ = QFileDialog.getOpenFileName(
                self,
                "Open DOCX file",
                "",
                "DOCX files (*.docx);;All files (*)",
                options=options
            )
            
            if filePath:
                if filePath.lower().endswith('.docx'):
                    self.convertFile(filePath)
                else:
                    QMessageBox.warning(self, "Invalid File", "Please select a .docx file.")
            else:
                print("No file selected.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {str(e)}")
            print(f"Error in openFileDialog: {str(e)}")

    def convertFile(self, filePath):
        try:
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            self.thread = ConversionThread(filePath, self.format_combo.currentText())
            self.thread.progress.connect(self.updateProgress)
            self.thread.finished.connect(self.conversionFinished)
            self.thread.error.connect(self.conversionError)
            self.thread.start()
        except Exception as e:
            self.conversionError(str(e))

    def updateProgress(self, value):
        self.progress_bar.setValue(value)

    def conversionFinished(self, output_path):
        self.progress_bar.setVisible(False)
        self.label.setText(f"Conversion successful: {output_path}")
        QMessageBox.information(self, "Success", f"File converted successfully:\n{output_path}")

    def conversionError(self, error_msg):
        self.progress_bar.setVisible(False)
        self.label.setText(f"Error: {error_msg}")
        QMessageBox.critical(self, "Conversion Error", f"An error occurred during conversion:\n{error_msg}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setFont(QFont("Segoe UI", 9))
    ex = DocxConverter()
    ex.show()
    sys.exit(app.exec())
